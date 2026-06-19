"""Servei d'extracció i anàlisi inicial de currículums PDF i DOCX.

Aquest mòdul no coneix res de la interfície web ni del servidor HTTP. La seva
única responsabilitat és rebre un document en memòria i retornar dades
estructurades. Aquesta separació permet substituir l'extractor per docTR, OCR o
un LLM en el futur sense haver de modificar la resta de l'aplicació.
"""

from __future__ import annotations

import io
import re
from collections import Counter
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_FILE_SIZE = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx"}

SECTION_WORDS = {
    "experiència": ("experiència", "experiencia", "experience", "employment", "work history"),
    "formació": ("formació", "formacion", "formación", "education", "estudis", "studies"),
    "habilitats": ("habilitats", "habilidades", "skills", "competències", "competencias"),
    "idiomes": ("idiomes", "idiomas", "languages"),
    "projectes": ("projectes", "proyectos", "projects"),
}

SKILLS = (
    "Python", "Java", "JavaScript", "TypeScript", "HTML", "CSS", "React", "Angular",
    "Vue", "Node.js", "FastAPI", "Django", "Flask", "Spring", "SQL", "MySQL",
    "PostgreSQL", "MongoDB", "Git", "GitHub", "Docker", "Kubernetes", "AWS",
    "Azure", "Excel", "Power BI", "Tableau", "Figma", "Photoshop", "Illustrator",
    "Scrum", "Agile", "Machine Learning", "Intel·ligència Artificial", "AI",
    "Data Analysis", "Anàlisi de dades", "Atenció al client", "Vendes", "Marketing",
    "SEO", "Comunicació", "Lideratge", "Treball en equip",
)

LANGUAGES = (
    "català", "catalan", "castellà", "español", "spanish", "anglès", "inglés",
    "english", "francès", "francés", "french", "alemany", "alemán", "german",
    "italià", "italiano", "portuguès", "portugués",
)

DEGREE_MARKERS = (
    "grau", "grado", "degree", "màster", "máster", "master", "batxillerat",
    "bachillerato", "formació professional", "formación profesional", "cicle",
    "ciclo formativo", "fp ", "universitat", "universidad", "university",
)

STOPWORDS = {
    "amb", "per", "una", "del", "dels", "les", "els", "que", "the", "and", "for",
    "with", "from", "this", "that", "para", "con", "por", "las", "los", "com",
    "més", "años", "anys", "year", "years", "sobre", "entre", "també", "tambien",
    "professional", "experiencia", "experiència", "experience", "formació",
    "formación", "education", "curriculum", "vitae",
}

TEXTS = {
    "ca": {
        "invalid_type": "Només s'accepten documents PDF o DOCX.",
        "too_large": "El document supera el límit de 10 MB.",
        "scanned": "No s'ha pogut extreure prou text. El PDF pot estar escanejat i necessitar OCR.",
        "short_docx": "El DOCX no conté prou text per analitzar-lo.",
        "invalid_pdf": "El PDF no és vàlid o està protegit.",
        "invalid_docx": "El DOCX no és vàlid o està malmès.",
        "skills": "Competències detectades: {value}.",
        "education": "Formació identificada: {value}",
        "timeline": "Experiència o cronologia detectada: {value}",
        "languages": "Idiomes mencionats: {value}.",
        "subject": "El perfil",
        "skill_detail": "destaca en {value}",
        "education_detail": "inclou formació acadèmica identificable",
        "timeline_detail": "presenta experiència o trajectòria amb dates",
        "summary": "{subject} és un perfil que {details}.",
        "summary_empty": "{subject} conté informació professional llegible, però caldrà aprofundir-hi amb preguntes.",
        "joiner": " i ",
        "fallback": "S'ha extret text del document, però no s'han identificat seccions clares.",
    },
    "es": {
        "invalid_type": "Solo se aceptan documentos PDF o DOCX.",
        "too_large": "El documento supera el límite de 10 MB.",
        "scanned": "No se ha podido extraer suficiente texto. El PDF puede estar escaneado y necesitar OCR.",
        "short_docx": "El DOCX no contiene suficiente texto para analizarlo.",
        "invalid_pdf": "El PDF no es válido o está protegido.",
        "invalid_docx": "El DOCX no es válido o está dañado.",
        "skills": "Competencias detectadas: {value}.",
        "education": "Formación identificada: {value}",
        "timeline": "Experiencia o cronología detectada: {value}",
        "languages": "Idiomas mencionados: {value}.",
        "subject": "El perfil",
        "skill_detail": "destaca en {value}",
        "education_detail": "incluye formación académica identificable",
        "timeline_detail": "presenta experiencia o trayectoria con fechas",
        "summary": "{subject} es un perfil que {details}.",
        "summary_empty": "{subject} contiene información profesional legible, pero será necesario profundizar mediante preguntas.",
        "joiner": " y ",
        "fallback": "Se ha extraído texto del documento, pero no se han identificado secciones claras.",
    },
    "en": {
        "invalid_type": "Only PDF or DOCX documents are accepted.",
        "too_large": "The document exceeds the 10 MB limit.",
        "scanned": "Not enough text could be extracted. The PDF may be scanned and require OCR.",
        "short_docx": "The DOCX does not contain enough text to analyze.",
        "invalid_pdf": "The PDF is invalid or protected.",
        "invalid_docx": "The DOCX is invalid or damaged.",
        "skills": "Skills detected: {value}.",
        "education": "Education identified: {value}",
        "timeline": "Experience or timeline detected: {value}",
        "languages": "Languages mentioned: {value}.",
        "subject": "The candidate",
        "skill_detail": "stands out in {value}",
        "education_detail": "includes identifiable academic education",
        "timeline_detail": "has experience or a career timeline with dates",
        "summary": "{subject} is a profile that {details}.",
        "summary_empty": "{subject} contains readable professional information, but follow-up questions will be needed.",
        "joiner": " and ",
        "fallback": "Text was extracted from the document, but no clear sections were identified.",
    },
}


def extract_text(filename: str, content: bytes, language: str = "ca") -> tuple[str, dict]:
    language = language if language in TEXTS else "ca"
    labels = TEXTS[language]
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError(labels["invalid_type"])
    if len(content) > MAX_FILE_SIZE:
        raise ValueError(labels["too_large"])

    if extension == ".pdf":
        text, pages = _extract_pdf(content, labels)
        metadata = {"format": "PDF", "pages": pages}
    else:
        text, paragraphs = _extract_docx(content, labels)
        metadata = {"format": "DOCX", "paragraphs": paragraphs}

    text = _clean_text(text)
    if len(text) < 40:
        if extension == ".pdf":
            raise ValueError(labels["scanned"])
        raise ValueError(labels["short_docx"])

    metadata["characters"] = len(text)
    metadata["words"] = len(text.split())
    return text, metadata


def _extract_pdf(content: bytes, labels: dict) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        raise ValueError(labels["invalid_pdf"]) from exc
    return "\n".join(pages), len(pages)


def _extract_docx(content: bytes, labels: dict) -> tuple[str, int]:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(labels["invalid_docx"]) from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))

    return "\n".join(parts), len(parts)


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\u00ad", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def analyse_cv(filename: str, content: bytes, language: str = "ca") -> dict:
    language = language if language in TEXTS else "ca"
    labels = TEXTS[language]
    text, metadata = extract_text(filename, content, language)
    lines = [line.strip(" •·-\t") for line in text.splitlines() if line.strip()]

    emails = _unique(re.findall(r"[\w.+-]+@[\w-]+\.[\w.-]+", text))
    phones = _unique(re.findall(r"(?:\+\d{1,3}[\s.-]?)?(?:\d[\s.-]?){8,12}", text))
    links = _unique(re.findall(r"(?:https?://|www\.)\S+|linkedin\.com/\S+", text, re.I))
    skills = [skill for skill in SKILLS if _contains_term(text, skill)]
    languages = [language.title() for language in LANGUAGES if _contains_term(text, language)]
    education = _matching_lines(lines, DEGREE_MARKERS, limit=4)
    dated_lines = [
        line for line in lines
        if re.search(r"\b(?:19|20)\d{2}\b", line) or re.search(r"\b\d{1,2}/\d{4}\b", line)
    ][:6]
    sections = [
        name for name, markers in SECTION_WORDS.items()
        if any(_contains_term(text, marker) for marker in markers)
    ]

    highlights = []
    if skills:
        highlights.append(labels["skills"].format(value=", ".join(skills[:10])))
    if education:
        highlights.append(labels["education"].format(value=education[0]))
    if dated_lines:
        highlights.append(labels["timeline"].format(value=dated_lines[0]))
    if languages:
        highlights.append(labels["languages"].format(value=", ".join(languages[:6])))
    if not highlights:
        highlights = _fallback_highlights(lines, labels)

    candidate_name = _guess_name(lines)
    profile_sentence = _profile_summary(candidate_name, skills, education, dated_lines, labels)
    keywords = _top_keywords(text, limit=8)

    return {
        "filename": filename,
        "summary": profile_sentence,
        "highlights": highlights[:5],
        "skills": skills[:16],
        "languages": languages[:8],
        "education": education,
        "timeline": dated_lines,
        "sections": sections,
        "keywords": keywords,
        "contact_detected": {
            "email": bool(emails),
            "phone": bool(phones),
            "links": len(links),
        },
        "metadata": metadata,
        "preview": text[:1200],
    }


def _contains_term(text: str, term: str) -> bool:
    return bool(re.search(rf"(?<!\w){re.escape(term)}(?!\w)", text, re.I))


def _matching_lines(lines: list[str], markers: tuple[str, ...], limit: int) -> list[str]:
    matches = []
    for line in lines:
        if 5 < len(line) < 180 and any(_contains_term(line, marker) for marker in markers):
            matches.append(line)
    return _unique(matches)[:limit]


def _guess_name(lines: list[str]) -> str:
    for line in lines[:8]:
        if (
            1 < len(line.split()) <= 5
            and len(line) < 60
            and not re.search(r"[@:/\d]", line)
            and not any(word in line.lower() for word in ("curriculum", "currículum", "resume", "perfil"))
        ):
            return line.title() if line.isupper() else line
    return ""


def _profile_summary(
    name: str, skills: list[str], education: list[str], dated_lines: list[str], labels: dict
) -> str:
    subject = name or labels["subject"]
    details = []
    if skills:
        details.append(labels["skill_detail"].format(value=", ".join(skills[:4])))
    if education:
        details.append(labels["education_detail"])
    if dated_lines:
        details.append(labels["timeline_detail"])

    if details:
        return labels["summary"].format(subject=subject, details=labels["joiner"].join(details))
    return labels["summary_empty"].format(subject=subject)


def _fallback_highlights(lines: list[str], labels: dict) -> list[str]:
    useful = [
        line for line in lines
        if 35 <= len(line) <= 180 and not re.search(r"@|https?://|www\.", line, re.I)
    ]
    return useful[:4] or [labels["fallback"]]


def _top_keywords(text: str, limit: int) -> list[str]:
    words = re.findall(r"[A-Za-zÀ-ÿ][A-Za-zÀ-ÿ+#.]{2,}", text.lower())
    counts = Counter(word for word in words if word not in STOPWORDS and len(word) > 3)
    return [word for word, _ in counts.most_common(limit)]


def _unique(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        normalized = re.sub(r"\s+", " ", item).strip()
        key = normalized.casefold()
        if normalized and key not in seen:
            seen.add(key)
            result.append(normalized)
    return result
